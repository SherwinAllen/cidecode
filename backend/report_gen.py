from pymongo import MongoClient
import gridfs
import pandas as pd
import docx
import re
from collections import Counter
import tempfile
import matplotlib.pyplot as plt
from fastapi.responses import FileResponse
from flask import Flask
import datetime
import os
import json

app = Flask(__name__)

# ---------------- MongoDB Setup ----------------
client = MongoClient("mongodb://localhost:27017/")
db = client["forensic_evidence"]
fs = gridfs.GridFS(db)


def get_file_from_mongo(filename):
    """Fetch a file from MongoDB GridFS and return content as text."""
    file_doc = fs.find_one({"filename": filename})
    if not file_doc:
        print(f"[-] File '{filename}' not found in MongoDB.")
        return ""
    try:
        print(f"File {filename} Found")
        data = file_doc.read()
        # Decode text files; binary files can be handled separately if needed
        return data.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"[!] Error reading {filename} from MongoDB: {e}")
        return ""

def extract_logs_from_file(filepath):
    """Reads up to 20 lines from the given file."""
    parsed_data = []
    raw_lines = []
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            lines = file.readlines()
            for line in lines[:20]:
                raw_lines.append(line.rstrip('\n'))
                parsed_data.append(line.strip().split())
    except PermissionError as e:
        print(f"Permission denied when accessing {filepath}: {e}")
    return parsed_data, raw_lines


def parse_sensor_line(line):
    """Extracts structured info for each sensor log line."""
    pattern = re.compile(
        r'(?P<sensor_id>0x[0-9a-f]+)\).*active-count\s*=\s*(?P<active_count>\d+);'
        r'.*sampling_period\(ms\)\s*=\s*\{(?P<sampling>[\d., ]+)\}.*'
        r'batching_period\(ms\)\s*=\s*\{(?P<batching>[\d., ]+)\}.*selected\s*=\s*(?P<selected>[\d.]+) ms'
    )
    match = pattern.search(line)
    if match:
        return [
            match.group("sensor_id"),
            match.group("active_count"),
            match.group("sampling").replace(",", ", "),
            match.group("batching").replace(",", ", "),
            match.group("selected")
        ]
    else:
        return [""] * 5


def parse_account_info(log_text):
    """
    Parses Android account-related forensic dump text into two DataFrames:
      - Accounts table
      - Registered Services table
    """
    accounts = []
    services = []

    # --- Parse accounts ---
    for line in log_text.splitlines():
        line = line.strip()
        # Match Account {name=..., type=...}
        acc_match = re.match(r'Account\s*\{name=([^,]+),\s*type=([^}]+)\}', line)
        if acc_match:
            name = acc_match.group(1).strip()
            acc_type = acc_match.group(2).strip()
            accounts.append({"Account Name": name, "Type": acc_type})
            continue

        # Match ServiceInfo lines
        svc_match = re.match(
            r'ServiceInfo:\s*AuthenticatorDescription\s*\{type=([^}]+)\},\s*ComponentInfo\{([^}]+)\},\s*uid\s*(\d+)',
            line
        )
        if svc_match:
            svc_type = svc_match.group(1).strip()
            component = svc_match.group(2).strip()
            uid = svc_match.group(3).strip()
            services.append({
                "Type": svc_type,
                "Component": component,
                "UID": uid
            })

    # Convert to DataFrames
    accounts_df = pd.DataFrame(accounts)
    services_df = pd.DataFrame(services)

    return accounts_df, services_df


def extract_sensor_timestamps(raw_lines):
    """
    Extract timestamps, sensor IDs, and (optional) sensor names.
    Returns dict: {sensor_id: DataFrame}
    """
    sensor_data = {}
    # Regex captures timestamps, sensor IDs, and possible sensor names
    pattern = re.compile(
        r'(?P<timestamp>\d{2}-\d{2}\s\d{2}:\d{2}:\d{2}\.\d+).*?'
        r'(?P<sensor_id>0x[0-9a-f]+).*?'
        r'(?:SensorName\s*=\s*(?P<name>[A-Za-z0-9_ -]+))?',
        re.IGNORECASE
    )

    for line in raw_lines:
        match = pattern.search(line)
        print(line)
        if match:
            timestamp = match.group("timestamp")
            sensor_id = match.group("sensor_id")
            sensor_name = match.group("name") or "Unknown Sensor"
            if sensor_id not in sensor_data:
                sensor_data[sensor_id] = {"name": sensor_name, "entries": []}
            sensor_data[sensor_id]["entries"].append((timestamp, line.strip()))

    # Convert to DataFrames
    sensor_dfs = {}
    for sensor_id, info in sensor_data.items():
        df = pd.DataFrame(info["entries"], columns=["Timestamp", "Log Line"])
        sensor_dfs[sensor_id] = (info["name"], df)
    return sensor_dfs


def add_dataframe_to_doc(doc, df, title, max_cols_per_table=5):
    """
    Writes a pandas DataFrame into the Word doc as formatted tables.
    Splits wide DataFrames into multiple tables if columns exceed max_cols_per_table.
    """
    if df.empty:
        doc.add_paragraph(f"{title} - No data found.\n", style='Heading3')
        return

    columns = df.columns.tolist()
    start = 0
    table_index = 1

    while start < len(columns):
        subset_cols = columns[start:start + max_cols_per_table]
        sub_df = df[subset_cols]

        # Add title for this part
        if len(columns) > max_cols_per_table:
            doc.add_paragraph(f"{title} (Part {table_index})", style='Heading4')
        else:
            doc.add_paragraph(title, style='Heading3')

        # Create table
        table = doc.add_table(rows=1, cols=len(sub_df.columns))
        table.style = "Table Grid"

        # Header row
        for i, col_name in enumerate(sub_df.columns):
            table.cell(0, i).text = col_name
            for run in table.cell(0, i).paragraphs[0].runs:
                run.bold = True

        # Data rows
        for _, row in sub_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        doc.add_paragraph("\n")
        start += max_cols_per_table
        table_index += 1

# --- Configuration ---

log_files = {
    "Account Information": "account_information.txt",
    "Bluetooth Information": "bluetooth_information.txt",
    "Device Properties": "device_properties.txt",
    "Sensor Data": "sensor_data.txt",
    "Ip information": "ip_address_information.txt"
}

column_headers = {
    "Account Information": ["Field", "Value"],
    "Bluetooth Information": ["Field", "Value"],
    "Device Properties": ["Field", "Value"],
    "Sensor Data": ["Timestamp", "Sensor Type", "Value"]
}



# =============================================================== DEBUG =================================================================
def extract_sensor_data(log_text):
    sensors = {}
    current_sensor = None
    records = []

    for line in log_text.splitlines():
        line = line.strip()

        # --- Detect start of a new sensor section ---
        match_sensor = re.match(r'^(.*?):.*events$', line)
        print(line)
        if match_sensor:
            # Save previous sensor’s records
            if current_sensor and records:
                sensors[current_sensor] = pd.DataFrame(records)
                records = []
            current_sensor = match_sensor.group(1).strip()
            continue

        # --- Match event lines ---
        # Flexible regex for formats like:
        # 1 (ts=123.456, wall=12:34:56.789) 1.00, 0.00,
        match_event = re.match(
            r'^\d+\s*\(ts=([\d.]+),\s*wall=([\d:.]+)\)\s*(.*)', line
        )
        if match_event:
            ts = float(match_event.group(1))
            wall = match_event.group(2).strip()
            values_str = match_event.group(3).strip()

            # --- Handle value formats ---
            if "[value masked]" in values_str:
                record = {"ts": ts, "wall_time": wall, "values": "[value masked]"}
            elif values_str:
                # Extract all numeric values, even if followed by commas
                nums = [float(x) for x in re.findall(r'[-+]?\d*\.\d+|\d+', values_str)]
                record = {"ts": ts, "wall_time": wall}
                for i, val in enumerate(nums, start=1):
                    record[f"value_{i}"] = val
            else:
                record = {"ts": ts, "wall_time": wall}

            records.append(record)

    # Save last sensor’s data
    if current_sensor and records:
        sensors[current_sensor] = pd.DataFrame(records)

    return sensors


def parse_bluetooth_log(doc, text):
    """Extracts Bluetooth connection and bonded device info."""

    # 1️⃣ Connection / Disconnection events
    print(text)
    dates = re.findall(r"(\d{2}-\d{2})\s\d{2}:\d{2}:\d{2}\.\d{3}", text)

    print(dates)
    # Count number of events per day
    counter = Counter(dates)

    # Convert to pandas Series for plotting
    df = pd.Series(counter).sort_index()
    title = "Bluetooth Events Per Day"
    # Plot bar graph
    plt.figure(figsize=(10,6))
    df.plot(kind='bar', color='skyblue')
    plt.xlabel("Date (MM - DD)")
    plt.ylabel("Number of Bluetooth Events")
    plt.title("Bluetooth Events per Day")
    plt.xticks(rotation=45)

    # Save plot to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        plt.savefig(tmpfile.name, format='png')
        plt.close()
        doc.add_paragraph(title, style='Heading3')
        doc.add_picture(tmpfile.name, width=docx.shared.Inches(6))

    # 2️⃣ Bonded devices
    bonded_pattern = r"\s*\(Connected\)\s*([0-9A-F:]{17}) \[.*?\] ([^\(]+)"

    bonded_match = re.findall(bonded_pattern, text)
    bonded_devices = []
    if bonded_match:
        for match in bonded_match:
            bonded_devices.append({
                    "Device Name": match[1],
                    "MAC Address": match[0]
                })

    df_bonded = pd.DataFrame(bonded_devices)

    return df_bonded


def extract_ip_info(output_text):
    """
    Parse ADB `ip addr` output and return a pandas DataFrame
    with columns: Interface, Status, MTU, IPv4, Broadcast, IPv6, MAC, Notes
    mapped to value_1, value_2, ..., value_8 for consistency.
    """
    interfaces = []
    current_iface = {}
    for line in output_text.splitlines():
        line = line.strip()
        if re.match(r'^\d+:', line):
            # New interface line
            if current_iface:
                interfaces.append(current_iface)
            m = re.match(r'^(\d+):\s+([\w@]+):\s+<([^>]*)>.*mtu\s+(\d+)', line)
            if m:
                current_iface = {
                    "Interface": m.group(2),
                    "Status": "UP" if "UP" in m.group(3).split(",") else "DOWN",
                    "MTU": int(m.group(4)),
                    "IPv4": "",
                    "Broadcast": "",
                    "IPv6": "",
                    "MAC": "",
                    "Notes": ""
                }
        elif line.startswith("link/"):
            # MAC address line
            m = re.match(r'link/\w+\s+([\da-f:]+)', line)
            if m:
                current_iface["MAC"] = m.group(1)
        elif line.startswith("inet "):
            # IPv4 line
            m = re.match(r'inet\s+([\d./]+)\s+brd\s+([\d.]+)', line)
            if m:
                current_iface["IPv4"] = m.group(1)
                current_iface["Broadcast"] = m.group(2)
        elif line.startswith("inet6 "):
            # IPv6 line
            m = re.match(r'inet6\s+([\da-f:]+/[\d]+)', line)
            if m:
                current_iface["IPv6"] = m.group(1)

    if current_iface:
        interfaces.append(current_iface)

    # Convert to DataFrame and map to value_1..value_8
    df = pd.DataFrame(interfaces)
    # df = df.rename(columns={
    #     "Interface": "value_1",
    #     "Status": "value_2",
    #     "MTU": "value_3",
    #     "IPv4": "value_4",
    #     "Broadcast": "value_5",
    #     "IPv6": "value_6",
    #     "MAC": "value_7",
    #     "Notes": "value_8"
    # })
    return df

def get_location(location_file):
    with open(location_file,'r', encoding = 'utf-8') as f:
        location_text = f.read()
        regex = re.compile(
            r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?(?:hAcc=(?P<acc>\d+\.?\d*))?',
            re.IGNORECASE | re.DOTALL
        )

        matches = list(regex.finditer(location_text))
        records = []
        for m in matches:
            records.append({
                "provider": m.group("provider"),
                "latitude": float(m.group("lat")),
                "longitude": float(m.group("lon")),
                "accuracy": float(m.group("acc")) if m.group("acc") else None
            })

        # Convert to DataFrame
        df = pd.DataFrame(records)

    return df


def parse_wifi_log_extended(log_text: str):
    """
    Parse ADB Wi-Fi diagnostic logs including:
      - SSID/BSSID connection info
      - Connection metrics
      - Supplicant state transitions
      - Multi-Link (Mlink) info
    Returns a dict of DataFrames.
    """
    dfs = {}

    # 1️⃣ ---- Wi-Fi SSID/BSSID Info ----
    ssid_pattern = re.compile(
        r'rec\[\d+\]:\s+'                      # rec number
        r'time=(?P<timestamp>[\d\-:\. ]+)\s+'  # timestamp
        r'processed=(?P<processed>\S+)\s+'
        r'org=(?P<org>\S+)\s+'
        r'dest=(?P<dest>\S+)\s+'
        r'what=(?P<what>\S+)\s+'
        r'screen=\S+\s+\d+\s+\d+\s+'
        r'ssid:\s*"(?P<ssid>[^"]+)"\s+'
        r'bssid:\s*(?P<bssid>[0-9a-f:]+)\s+'
        r'nid:\s*(?P<nid>\d+)\s+'
        r'frequencyMhz:\s*(?P<freq>\d+)\s+'
        r'state:\s*COMPLETED', 
        re.IGNORECASE
    )
    ssid_records = []
    for line in log_text.splitlines():
        m = ssid_pattern.search(line)
        if m:
            ssid_records.append({
                "timestamp": m.group('timestamp'),
                "ssid": m.group('ssid'),
                "bssid": m.group('bssid')
            })

    # 2️⃣ ---- Wi-Fi Metrics ----
    wifi_pattern = re.compile(
        r"time=(?P<time>[\d\-\s:]+).*?"
        r"session=(?P<session>[^,]+),?"
        r".*?netid=(?P<netid>[^,]+),?"
        r".*?rssi=(?P<rssi>[^,]+),?"
        r".*?filtered_rssi=(?P<filtered_rssi>[^,]+),?"
        r".*?freq=(?P<freq>[^,]+),?"
        r".*?txLinkSpeed=(?P<txLinkSpeed>[^,]+),?"
        r".*?rxLinkSpeed=(?P<rxLinkSpeed>[^,]+),?",
        re.DOTALL
    )
    wifi_records = []
    for line in log_text.splitlines():
        if "rssi=" in line and "txLinkSpeed=" in line:
            m = wifi_pattern.search(line)
            if m:
                wifi_records.append(m.groupdict())

    # 3️⃣ ---- Supplicant State Tracker ----
    supplicant_pattern = re.compile(
        r"rec\[\d+\]: time=(?P<time>[\d\-:\.\s]+).*?"
        r"org=(?P<org_state>\S+).*?"
        r"dest=(?P<dest_state>\S*).*?"  # allow empty dest
        r"what=(?P<what>[0-9xXA-F]+)",
        re.DOTALL
    )
    supplicant_records = [m.groupdict() for m in supplicant_pattern.finditer(log_text)]

    # Filter out rows where dest_state is empty
    supplicant_records = [r for r in supplicant_records if r["dest_state"].strip() != "<null>"]

    # Convert to DataFrame
    if supplicant_records:
        dfs["supplicant_states"] = pd.DataFrame(supplicant_records)

    # 4️⃣ ---- Mlink / Multi-Link Operation ----
    mlink_pattern = re.compile(
        r"\{linkId=(?P<linkId>\d+),linkRssi=(?P<linkRssi>[^,]+),linkFreq=(?P<linkFreq>[^,]+),"
        r"txLinkSpeed=(?P<txLinkSpeed>[^,]+),rxLinkSpeed=(?P<rxLinkSpeed>[^,]+).*?\}",
        re.DOTALL
    )



    mlink_records = [m.groupdict() for m in mlink_pattern.finditer(log_text)]

    # Convert to DataFrames
    if ssid_records:
        dfs["wifi_networks"] = pd.DataFrame(ssid_records)
    if wifi_records:
        dfs["wifi_metrics"] = pd.DataFrame(wifi_records)
    if supplicant_records:
        dfs["supplicant_states"] = pd.DataFrame(supplicant_records)
    if mlink_records:
        dfs["mlink_info"] = pd.DataFrame(mlink_records)

    return dfs


def parse_location_data(loc_path):
    with open(loc_path, 'r', encoding="utf-8") as f:
        log_text = f.read()
    
    regex = re.compile(
        r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*'
        r'(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?'
        r'(?:hAcc=(?P<acc>\d+\.?\d*))?',
        re.IGNORECASE | re.DOTALL
    )

    matches = list(regex.finditer(log_text))
    if not matches:
        print("[-] No coordinate patterns found in dumpsys output.")
        df = pd.DataFrame(columns=["timestamp", "provider", "lat", "lon", "accuracy"])
    else:
        rows = []
        for m in matches:
            provider = m.group('provider') or 'unknown'
            lat = float(m.group('lat'))
            lon = float(m.group('lon'))
            acc = m.group('acc')
            ts = datetime.datetime.now().isoformat()
            rows.append({
                "timestamp": ts,
                "provider": provider,
                "lat": lat,
                "lon": lon,
                "accuracy": float(acc) if acc else None
            })

        df = pd.DataFrame(rows)
        print(f"✅ Parsed {len(df)} location entries.")
        print(df)

log_files = {
    "Account Information": "account_information.txt",
    "Bluetooth Information": "bluetooth_information.txt",
    "Device Properties": "device_properties.txt",
    "Sensor Data": "sensor_data.txt",
    "Ip information": "ip_address_information.txt"
}

column_headers = {
    "Account Information": ["Field", "Value"],
    "Bluetooth Information": ["Field", "Value"],
    "Device Properties": ["Field", "Value"],
    "Sensor Data": ["Timestamp", "Sensor Type", "Value"]
}

# ----------------------------------------------------------------
# Function that builds the forensic report and saves it
# ----------------------------------------------------------------
log_files = {
    "Account Information": "account_information.txt",
    "Bluetooth Information": "bluetooth_information.txt",
    "Device Properties": "device_properties.txt",
    "Sensor Data": "sensor_data.txt",
    "Ip information": "ip_address_information.txt",
    "WiFi Information": "wifi_information.txt",
    "Location Information": "dumpsys_location.txt"
}

# ---------------- Forensic Report Generation ----------------
def generate_forensic_report(output_dir="downloads"):
    """Generates the forensic .docx report using MongoDB data."""
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Forensic_Log_Report.docx")

    doc = docx.Document()
    doc.add_paragraph("Forensic Log Report", style='Title')
    
    
    # --- Account Info ---
    acc_text = get_file_from_mongo(log_files["Account Information"])
    acc_df, service_df = parse_account_info(acc_text)
    add_dataframe_to_doc(doc, acc_df, "Account Information")
    add_dataframe_to_doc(doc, service_df, "Service Information")

    # --- Wi-Fi Info ---
    wifi_text = get_file_from_mongo(log_files["WiFi Information"])
    wifi_df_dict = parse_wifi_log_extended(wifi_text)
    for section_name, df in wifi_df_dict.items():
        add_dataframe_to_doc(doc, df, f"Wi-Fi: {section_name.replace('_', ' ').title()}")

    # --- Bluetooth Info ---
    bt_text = get_file_from_mongo(log_files["Bluetooth Information"])
    df_bonded = parse_bluetooth_log(doc, bt_text)
    add_dataframe_to_doc(doc, df_bonded, "Bonded Bluetooth Devices")

    # --- Location Info ---
    loc_text = get_file_from_mongo(log_files["Location Information"])
    loc_df = get_location_text(loc_text)
    add_dataframe_to_doc(doc, loc_df, "Location Information")

    # --- Sensor Data ---
    sensor_text = get_file_from_mongo(log_files["Sensor Data"])
    sensor_dataframes = extract_sensor_data(sensor_text)
    for sensor_name, df in sensor_dataframes.items():
        add_dataframe_to_doc(doc, df, sensor_name)

    # --- IP Info ---
    ip_text = get_file_from_mongo(log_files["Ip information"])
    ip_df = extract_ip_info(ip_text)
    add_dataframe_to_doc(doc, ip_df, "IP Address Information")

    # Save to DOCX
    doc.save(output_path)
    print(f"Forensic report saved to: {output_path}")
    return output_path

# ---------------- Flask Route ----------------
@app.get("/download_report")
def download_report():
    report_path = "downloads/Preliminary_Forensic_Report.docx"
    if os.path.exists(report_path):
        return FileResponse(
            path=report_path,
            filename="Preliminary_Forensic_report.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return {"error": "Report not found"}


# ---------------- Helper for location text ----------------
def get_location_text(location_text):
    """Parse location text from MongoDB (previously from file)."""
    regex = re.compile(
        r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*'
        r'(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?'
        r'(?:hAcc=(?P<acc>\d+\.?\d*))?',
        re.IGNORECASE | re.DOTALL
    )
    matches = list(regex.finditer(location_text))
    records = []
    for m in matches:
        records.append({
            "provider": m.group("provider") or "unknown",
            "latitude": float(m.group("lat")),
            "longitude": float(m.group("lon")),
            "accuracy": float(m.group("acc")) if m.group("acc") else None
        })
    return pd.DataFrame(records)

# ---------------- Main ----------------
if __name__ == "__main__":
    generate_forensic_report()